"""
Backend de extracción de transcripts — versión web (Render + MongoDB).

Reutiliza exactamente la misma lógica de extracción que ya validamos
en Colab (extractor_v4_playwright.js: scroll cuidadoso, orden real por
conversation-turn-N, texto estable, content_blocks para listas/tablas).

Diferencia respecto a la versión anterior: el registro estructurado de
cada conversación (metadatos + turnos completos) se guarda en MongoDB
en vez de Google Sheets/Drive -- evita por completo la fricción de
cuentas de servicio, cuotas de almacenamiento y permisos de Unidades
compartidas. El .docx para el profesor se sigue generando al vuelo y
entregando en un .zip, sin persistirse en ningún lado.
"""

import io
import os
import re
import time
import zipfile
from datetime import datetime, timezone
from typing import List, Optional

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from pymongo import MongoClient
import certifi

from docx import Document

from playwright.async_api import async_playwright


# ============================================================
# Configuración / conexión a MongoDB
# ============================================================

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"
)

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(_THIS_DIR, "extractor_v4_playwright.js"), "r", encoding="utf-8") as f:
    EXTRACTOR_JS = f.read()

COMBINED_SCRIPT = EXTRACTOR_JS + "\nreturn await window.__extractConversation();"

_mongo_client = None


def get_mongo_collection():
    """Conecta (una sola vez, reutilizando la conexión) a la colección
    donde se guarda un documento por conversación. La cadena de
    conexión completa (con usuario/contraseña) vive en la variable de
    entorno MONGODB_URI -- nunca en el código."""
    global _mongo_client
    uri = os.environ.get("MONGODB_URI")
    if not uri:
        raise RuntimeError(
            "Falta la variable de entorno MONGODB_URI "
            "(cadena de conexión completa de tu clúster de MongoDB Atlas)."
        )
    if _mongo_client is None:
        _mongo_client = MongoClient(uri, tlsCAFile=certifi.where())
    db_name = os.environ.get("MONGODB_DB_NAME", "transcripts")
    coleccion_name = os.environ.get("MONGODB_COLLECTION_NAME", "conversaciones")
    return _mongo_client[db_name][coleccion_name]


# ============================================================
# Modelos de la petición
# ============================================================

class Alumno(BaseModel):
    nombre: str
    link: str
    fecha_envio: Optional[str] = ""


class ActividadRequest(BaseModel):
    curso: str
    clave_materia: str
    programa: str
    arranque: str
    actividad_nombre: str
    actividad_numero: str
    fecha_limite: str
    alumnos: List[Alumno]


# ============================================================
# Utilidades
# ============================================================

def limpiar_link(link: str) -> str:
    link = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", link or "")
    return link.strip()


def extraer_resumen_adjuntos(resultado: dict) -> list:
    turns = resultado.get("turns", [])
    nombres_vistos = []
    for t in turns:
        for a in t.get("attachments", []):
            nombre = (a.get("name") or "").strip()
            if nombre and nombre not in nombres_vistos:
                nombres_vistos.append(nombre)
        for f in t.get("files_mentioned", []):
            nombre = (f.get("name") or "").strip()
            if nombre and nombre not in nombres_vistos:
                nombres_vistos.append(nombre)
    return nombres_vistos


async def extraer_una_conversacion(page, link: str) -> dict:
    await page.goto(link, wait_until="networkidle", timeout=45000)
    await page.wait_for_timeout(1500)
    return await page.evaluate("async () => {\n" + COMBINED_SCRIPT + "\n}")


def guardar_registro(coleccion, alumno: Alumno, req: ActividadRequest, resultado: dict = None, error: str = None):
    """Un documento por conversación (o por intento fallido), con los
    metadatos de la actividad + el resultado completo del extractor
    (todos los turnos, adjuntos, quality_report) tal cual -- sin
    aplanar nada, MongoDB guarda el JSON anidado directamente."""
    documento = {
        "nombre_alumno": alumno.nombre,
        "link": alumno.link,
        "fecha_envio": alumno.fecha_envio or "",
        "curso": req.curso,
        "clave_materia": req.clave_materia,
        "programa": req.programa,
        "arranque": req.arranque,
        "actividad_nombre": req.actividad_nombre,
        "actividad_numero": req.actividad_numero,
        "fecha_limite": req.fecha_limite,
        "procesado_en": datetime.now(timezone.utc).isoformat(),
    }

    if resultado is not None:
        conv_meta = resultado.get("conversation_metadata", {})
        documento.update({
            "estado": "ok",
            "conversation_id": conv_meta.get("conversation_id", ""),
            "total_turnos": conv_meta.get("total_turns", 0),
            "adjuntos": extraer_resumen_adjuntos(resultado),
            "resultado_completo": resultado,  # export_metadata, turns, analysis_units, quality_report
        })
    else:
        documento.update({
            "estado": "error",
            "error": error or "Error desconocido",
        })

    coleccion.insert_one(documento)


# ============================================================
# .docx: mismo contenido/estructura que ya validamos (listas,
# numeradas, tablas), sin cambios respecto a la versión anterior.
# ============================================================

def construir_docx(resultado: dict, alumno: Alumno, req: ActividadRequest) -> bytes:
    meta = resultado.get("export_metadata", {})
    conv_meta = resultado.get("conversation_metadata", {})
    turns = resultado.get("turns", [])

    doc = Document()

    doc.add_heading(meta.get("conversation_title") or "Transcript de conversación", level=0)

    info = doc.add_paragraph()
    info.add_run(
        f"URL: {meta.get('source_url', '')}\n"
        f"Extraído: {meta.get('exported_at', '')}\n"
        f"Turnos totales: {conv_meta.get('total_turns', 0)} "
        f"(usuario: {conv_meta.get('total_user_turns', 0)}, "
        f"asistente: {conv_meta.get('total_assistant_turns', 0)})"
    )

    procedencia = doc.add_paragraph()
    procedencia.add_run(
        f"Curso: {req.curso} ({req.clave_materia})\n"
        f"Programa: {req.programa}\n"
        f"Arranque: {req.arranque}\n"
        f"Actividad: {req.actividad_numero} - {req.actividad_nombre}\n"
        f"Fecha de envío del alumno: {alumno.fecha_envio or 'N/D'}\n"
        f"Fecha límite de la actividad: {req.fecha_limite}"
    )

    for t in turns:
        content_blocks = t.get("content_blocks")
        texto_plano = (t.get("text_clean") or "").strip()

        if not content_blocks and not texto_plano:
            continue

        etiqueta = f"{t.get('role_label', 'Desconocido')} — Turno {t.get('turn_index', '?')}"
        doc.add_heading(etiqueta, level=2)

        bloques = content_blocks if content_blocks else [{"type": "paragraph", "text": texto_plano}]

        for cb in bloques:
            tipo = cb.get("type")
            if tipo == "paragraph":
                texto = (cb.get("text") or "").strip()
                if texto:
                    doc.add_paragraph(texto)
            elif tipo == "bulleted_list":
                for item in cb.get("items", []):
                    item = item.strip()
                    if item:
                        doc.add_paragraph(item, style="List Bullet")
            elif tipo == "ordered_list":
                for item in cb.get("items", []):
                    item = item.strip()
                    if item:
                        doc.add_paragraph(item, style="List Number")
            elif tipo == "table":
                filas = [r for r in cb.get("rows", []) if r]
                if not filas:
                    continue
                celdas_por_fila = [f.split(" | ") for f in filas]
                num_cols = max(len(f) for f in celdas_por_fila)
                tabla = doc.add_table(rows=len(celdas_por_fila), cols=num_cols)
                tabla.style = "Light Grid Accent 1"
                for i, fila in enumerate(celdas_por_fila):
                    for j in range(num_cols):
                        tabla.cell(i, j).text = fila[j] if j < len(fila) else ""

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ============================================================
# App FastAPI
# ============================================================

app = FastAPI(title="Transcript Extractor Backend")

# El frontend en Vercel llama a este backend desde otro dominio, así
# que hay que permitirlo explícitamente. ALLOWED_ORIGIN se configura
# como variable de entorno en Render una vez que tengas la URL de
# Vercel (ej. "https://transcript-anahuac.vercel.app"). Mientras no
# esté configurada, se permite cualquier origen para no bloquearte
# durante las pruebas -- ajústalo antes de usarlo con datos reales.
_allowed_origin = os.environ.get("ALLOWED_ORIGIN", "*")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[_allowed_origin] if _allowed_origin != "*" else ["*"],
    allow_methods=["POST", "GET"],
    allow_headers=["*"],
)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/generate")
async def generate(req: ActividadRequest):
    if not req.alumnos:
        raise HTTPException(status_code=400, detail="La lista de alumnos está vacía.")

    try:
        coleccion = get_mongo_collection()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo conectar a la base de datos: {e}")

    zip_buffer = io.BytesIO()

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context(user_agent=USER_AGENT)
        page = await context.new_page()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for idx, alumno in enumerate(req.alumnos, start=1):
                link = limpiar_link(alumno.link)

                if not link:
                    guardar_registro(coleccion, alumno, req, error="Sin link")
                    continue

                try:
                    resultado = await extraer_una_conversacion(page, link)

                    if resultado.get("error"):
                        raise RuntimeError(resultado["error"])

                    total_turnos = resultado.get("conversation_metadata", {}).get("total_turns", 0)
                    if total_turnos == 0:
                        raise RuntimeError("El extractor no encontró turnos (0 mensajes detectados).")

                    docx_bytes = construir_docx(resultado, alumno, req)
                    nombre_archivo = f"Transcript - {alumno.nombre or f'alumno_{idx}'}.docx"
                    zf.writestr(nombre_archivo, docx_bytes)

                    guardar_registro(coleccion, alumno, req, resultado=resultado)

                except Exception as e:
                    guardar_registro(coleccion, alumno, req, error=str(e))

                time.sleep(1)  # cortesía hacia chatgpt.com/share

        await browser.close()

    zip_buffer.seek(0)
    nombre_zip = f"{req.actividad_numero} - {req.actividad_nombre} ({req.arranque})"
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{nombre_zip}.zip"'},
    )
