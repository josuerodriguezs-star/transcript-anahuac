"""
Backend de extracción de transcripts — versión web (Render).

Reutiliza exactamente la misma lógica que ya validamos en Colab:
  - extractor_v4_playwright.js (scroll cuidadoso, orden real por
    conversation-turn-N, texto estable, content_blocks para
    listas/tablas)
  - Subida del JSON crudo a Drive
  - Fila de metadatos en un Google Sheet (uno por actividad, dentro
    de la carpeta de Drive indicada)
  - Generación de un .docx por alumno con jerarquía real (Word usa
    estilos nativos, no símbolos de markdown)
  - Entrega final: un .zip con todos los .docx

Diferencia clave respecto a Colab: aquí NO se generan Google Docs.
"""

import io
import json
import os
import re
import time
import zipfile
from typing import List, Optional

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
import gspread

from docx import Document
from docx.shared import Pt

from playwright.async_api import async_playwright


# ============================================================
# Configuración / credenciales
# ============================================================

SCOPES = [
    "https://www.googleapis.com/auth/drive",
    "https://www.googleapis.com/auth/spreadsheets",
]

USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36"
)

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
with open(os.path.join(_THIS_DIR, "extractor_v4_playwright.js"), "r", encoding="utf-8") as f:
    EXTRACTOR_JS = f.read()

COMBINED_SCRIPT = EXTRACTOR_JS + "\nreturn await window.__extractConversation();"


def get_google_credentials():
    """Lee la clave JSON de la cuenta de servicio desde la variable de
    entorno GOOGLE_SERVICE_ACCOUNT_JSON (contenido completo del .json,
    no una ruta de archivo) -- así el secreto nunca vive en el código
    ni en el repositorio, solo en la configuración de Render."""
    raw = os.environ.get("GOOGLE_SERVICE_ACCOUNT_JSON")
    if not raw:
        raise RuntimeError(
            "Falta la variable de entorno GOOGLE_SERVICE_ACCOUNT_JSON "
            "(debe contener el JSON completo de la cuenta de servicio)."
        )
    info = json.loads(raw)
    return service_account.Credentials.from_service_account_info(info, scopes=SCOPES)


# ============================================================
# Modelos de la petición
# ============================================================

class Alumno(BaseModel):
    nombre: str
    link: str
    fecha_envio: Optional[str] = ""


class ActividadRequest(BaseModel):
    carpeta_drive_id: str          # carpeta de Drive donde vive/se crea el Sheet
    curso: str
    clave_materia: str
    programa: str
    arranque: str
    actividad_nombre: str
    actividad_numero: str
    fecha_limite: str
    alumnos: List[Alumno]


# ============================================================
# Utilidades (idénticas en espíritu a las del notebook de Colab)
# ============================================================

def limpiar_link(link: str) -> str:
    link = re.sub(r"[\u200b\u200c\u200d\ufeff]", "", link or "")
    return link.strip()


def extraer_resumen_adjuntos(resultado: dict) -> str:
    turns = resultado.get("turns", [])
    nombres_vistos = []
    for t in turns:
        for a in t.get("attachments", []):
            nombre = (a.get("name") or "").strip()
            if nombre and nombre not in nombres_vistos:
                nombres_vistos.append(nombre)
        for f in t.get("files_mentioned", []):
            nombre = (f.get("name") or "").strip()
            if nombre and nombre not in nombres_vistos:
                nombres_vistos.append(nombre)
    return ", ".join(nombres_vistos) if nombres_vistos else "Sin adjuntos"


async def extraer_una_conversacion(page, link: str) -> dict:
    await page.goto(link, wait_until="networkidle", timeout=45000)
    await page.wait_for_timeout(1500)
    return await page.evaluate("async () => {\n" + COMBINED_SCRIPT + "\n}")


# ============================================================
# Drive: subir JSON crudo y mover archivos a la carpeta correcta
# ============================================================

def mover_a_carpeta(drive_service, file_id: str, carpeta_id: str):
    """Movimiento correcto: hay que quitar los parents anteriores, no
    solo agregar el nuevo -- addParents solo no reubica el archivo."""
    info = drive_service.files().get(
        fileId=file_id, fields="parents", supportsAllDrives=True
    ).execute()
    parents_actuales = info.get("parents", [])
    remove_parents = ",".join(parents_actuales) if parents_actuales else None

    kwargs = {
        "fileId": file_id,
        "addParents": carpeta_id,
        "fields": "id, parents",
        "supportsAllDrives": True,
    }
    if remove_parents:
        kwargs["removeParents"] = remove_parents

    drive_service.files().update(**kwargs).execute()


def subir_json_crudo(drive_service, carpeta_id: str, nombre: str, resultado: dict) -> str:
    contenido = json.dumps(resultado, ensure_ascii=False, indent=2)
    media = MediaIoBaseUpload(
        io.BytesIO(contenido.encode("utf-8")), mimetype="application/json", resumable=False
    )
    metadata = {
        "name": f"Transcript_JSON - {nombre}.json",
        "mimeType": "application/json",
        "parents": [carpeta_id],
    }
    file = drive_service.files().create(
        body=metadata, media_body=media, fields="id, webViewLink", supportsAllDrives=True
    ).execute()
    return file.get("webViewLink", "")


# ============================================================
# Sheet: crear (si no existe) y escribir metadatos
# ============================================================

def obtener_o_crear_sheet(gc, drive_service, carpeta_id: str, nombre_sheet: str):
    """Busca un Sheet con ese nombre DENTRO de la carpeta indicada; si
    no existe, lo crea con los encabezados listos."""
    query = (
        f"name = '{nombre_sheet}' and '{carpeta_id}' in parents "
        "and mimeType = 'application/vnd.google-apps.spreadsheet' and trashed = false"
    )
    resp = drive_service.files().list(
        q=query, fields="files(id, name)", supportsAllDrives=True, includeItemsFromAllDrives=True
    ).execute()
    archivos = resp.get("files", [])

    if archivos:
        sheet_id = archivos[0]["id"]
        sh = gc.open_by_key(sheet_id)
        return sh.sheet1

    sh = gc.create(nombre_sheet)
    mover_a_carpeta(drive_service, sh.id, carpeta_id)
    ws = sh.sheet1
    ws.update("A1:L1", [[
        "Nombre", "Link", "Conversation ID", "JSON (Drive)",
        "Curso", "Clave de materia", "Programa", "Arranque",
        "Actividad (nombre)", "Actividad (número)", "Fecha límite",
        "Fecha de envío", "Archivos adjuntos", "Observaciones",
    ]])
    return ws


def escribir_fila(ws, fila_index: int, alumno: Alumno, req: ActividadRequest,
                   conversation_id: str, json_link: str, adjuntos: str, error: str = ""):
    ws.update(f"A{fila_index}:N{fila_index}", [[
        alumno.nombre,
        alumno.link,
        conversation_id,
        json_link,
        req.curso,
        req.clave_materia,
        req.programa,
        req.arranque,
        req.actividad_nombre,
        req.actividad_numero,
        req.fecha_limite,
        alumno.fecha_envio or "",
        adjuntos,
        error,
    ]])


# ============================================================
# .docx: mismo contenido/estructura que ya validamos (listas,
# numeradas, tablas), pero con python-docx en vez de la API de Docs.
# ============================================================

def construir_docx(resultado: dict, alumno: Alumno, req: ActividadRequest) -> bytes:
    meta = resultado.get("export_metadata", {})
    conv_meta = resultado.get("conversation_metadata", {})
    turns = resultado.get("turns", [])

    doc = Document()

    doc.add_heading(meta.get("conversation_title") or "Transcript de conversación", level=0)

    info = doc.add_paragraph()
    info.add_run(
        f"URL: {meta.get('source_url', '')}\n"
        f"Extraído: {meta.get('exported_at', '')}\n"
        f"Turnos totales: {conv_meta.get('total_turns', 0)} "
        f"(usuario: {conv_meta.get('total_user_turns', 0)}, "
        f"asistente: {conv_meta.get('total_assistant_turns', 0)})"
    )

    procedencia = doc.add_paragraph()
    procedencia.add_run(
        f"Curso: {req.curso} ({req.clave_materia})\n"
        f"Programa: {req.programa}\n"
        f"Arranque: {req.arranque}\n"
        f"Actividad: {req.actividad_numero} - {req.actividad_nombre}\n"
        f"Fecha de envío del alumno: {alumno.fecha_envio or 'N/D'}\n"
        f"Fecha límite de la actividad: {req.fecha_limite}"
    )

    for t in turns:
        content_blocks = t.get("content_blocks")
        texto_plano = (t.get("text_clean") or "").strip()

        if not content_blocks and not texto_plano:
            continue

        etiqueta = f"{t.get('role_label', 'Desconocido')} — Turno {t.get('turn_index', '?')}"
        doc.add_heading(etiqueta, level=2)

        bloques = content_blocks if content_blocks else [{"type": "paragraph", "text": texto_plano}]

        for cb in bloques:
            tipo = cb.get("type")
            if tipo == "paragraph":
                texto = (cb.get("text") or "").strip()
                if texto:
                    doc.add_paragraph(texto)
            elif tipo == "bulleted_list":
                for item in cb.get("items", []):
                    item = item.strip()
                    if item:
                        doc.add_paragraph(item, style="List Bullet")
            elif tipo == "ordered_list":
                for item in cb.get("items", []):
                    item = item.strip()
                    if item:
                        doc.add_paragraph(item, style="List Number")
            elif tipo == "table":
                filas = [r for r in cb.get("rows", []) if r]
                if not filas:
                    continue
                celdas_por_fila = [f.split(" | ") for f in filas]
                num_cols = max(len(f) for f in celdas_por_fila)
                tabla = doc.add_table(rows=len(celdas_por_fila), cols=num_cols)
                tabla.style = "Light Grid Accent 1"
                for i, fila in enumerate(celdas_por_fila):
                    for j in range(num_cols):
                        tabla.cell(i, j).text = fila[j] if j < len(fila) else ""

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ============================================================
# App FastAPI
# ============================================================

app = FastAPI(title="Transcript Extractor Backend")

# El frontend en Vercel llama a este backend desde otro dominio, así
# que hay que permitirlo explícitamente. ALLOWED_ORIGIN se configura
# como variable de entorno en Render una vez que tengas la URL de
# Vercel (ej. "https://transcript-anahuac.vercel.app"). Mientras no
# esté configurada, se permite cualquier origen para no bloquearte
# durante las pruebas -- ajústalo antes de usarlo con datos reales.
_allowed_origin = os.environ.get("ALLOWED_ORIGIN", "*")
app.add_middleware(
    CORSMiddleware,
    allow_origins=[_allowed_origin] if _allowed_origin != "*" else ["*"],
    allow_methods=["POST", "GET"],
    allow_headers=["*"],
)


@app.get("/health")
def health():
    return {"status": "ok"}


@app.post("/generate")
async def generate(req: ActividadRequest):
    if not req.alumnos:
        raise HTTPException(status_code=400, detail="La lista de alumnos está vacía.")

    creds = get_google_credentials()
    gc = gspread.authorize(creds)
    drive_service = build("drive", "v3", credentials=creds)

    nombre_sheet = f"{req.actividad_numero} - {req.actividad_nombre} ({req.arranque})"
    ws = obtener_o_crear_sheet(gc, drive_service, req.carpeta_drive_id, nombre_sheet)

    zip_buffer = io.BytesIO()

    async with async_playwright() as p:
        browser = await p.chromium.launch(headless=True)
        context = await browser.new_context(user_agent=USER_AGENT)
        page = await context.new_page()

        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
            for idx, alumno in enumerate(req.alumnos, start=2):
                link = limpiar_link(alumno.link)

                if not link:
                    escribir_fila(ws, idx, alumno, req, "", "", "", "Sin link")
                    continue

                try:
                    resultado = await extraer_una_conversacion(page, link)

                    if resultado.get("error"):
                        raise RuntimeError(resultado["error"])

                    total_turnos = resultado.get("conversation_metadata", {}).get("total_turns", 0)
                    if total_turnos == 0:
                        raise RuntimeError("El extractor no encontró turnos (0 mensajes detectados).")

                    conversation_id = resultado.get("conversation_metadata", {}).get("conversation_id", "")
                    json_link = subir_json_crudo(
                        drive_service, req.carpeta_drive_id, alumno.nombre or f"fila_{idx}", resultado
                    )
                    adjuntos = extraer_resumen_adjuntos(resultado)

                    docx_bytes = construir_docx(resultado, alumno, req)
                    nombre_archivo = f"Transcript - {alumno.nombre or f'fila_{idx}'}.docx"
                    zf.writestr(nombre_archivo, docx_bytes)

                    escribir_fila(ws, idx, alumno, req, conversation_id, json_link, adjuntos)

                except Exception as e:
                    escribir_fila(ws, idx, alumno, req, "", "", "", f"Error: {e}")

                time.sleep(1)  # cortesía hacia chatgpt.com/share, igual que en Colab

        await browser.close()

    zip_buffer.seek(0)
    return StreamingResponse(
        zip_buffer,
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{nombre_sheet}.zip"'},
    )
